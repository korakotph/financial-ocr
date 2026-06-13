#!/usr/bin/env python3
"""
update_docx_results.py
อ่านผลจาก report/eval_v3_*_summary.json แล้วอัปเดต financial-ocr-paper-v2.docx

- ตาราง 4.1: เพิ่มคอลัมน์ "การทดลองที่ 3 (prompt ปรับปรุง)" พร้อมตัวเลขใหม่
- ตาราง 4.6: เพิ่มแถว "การทดลองที่ 3" สรุปผล
- ตาราง 4.7: อัปเดตแถว document_number / document_date ถ้าดีขึ้น
"""

import json, sys, re, copy
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Missing: pip install python-docx"); sys.exit(1)

sys.stdout.reconfigure(encoding="utf-8")

BASE_DIR   = Path(__file__).parent
REPORT_DIR = BASE_DIR / "report"
DOCX_PATH  = BASE_DIR.parent / "financial-ocr-paper-v2.docx"

# ─── 1. หาไฟล์ผลล่าสุด ────────────────────────────────────────────────────────

summaries = sorted(REPORT_DIR.glob("eval_v3_*_summary.json"))
if not summaries:
    print("❌ ไม่พบ report/eval_v3_*_summary.json — กรุณารัน run_eval_v3.py ก่อน")
    sys.exit(1)

latest = json.loads(summaries[-1].read_text(encoding="utf-8"))
print(f"✓ ใช้ผลจาก: {summaries[-1].name}")

# รวมค่าเฉลี่ยจากทุก round
rounds = latest.get("rounds", [])
if not rounds:
    print("❌ ไม่มีข้อมูล rounds ในไฟล์ summary"); sys.exit(1)

def avg_field(field_key):
    vals = [r["report"]["fields"].get(field_key, {}).get("rate", 0) for r in rounds]
    return round(sum(vals) / len(vals), 4) if vals else 0

# เก็บค่าสำคัญ
n_total  = rounds[0]["report"].get("n_total", 204)
n_succ   = round(sum(r["report"]["n_success"] for r in rounds) / len(rounds))
succ_rate = round(n_succ / n_total, 4)

fields_new = {
    "document_type":   avg_field("document_type"),
    "document_number": avg_field("document_number"),
    "document_date":   avg_field("document_date"),
    "seller.name":     avg_field("seller.name"),
    "seller.tax_id":   avg_field("seller.tax_id"),
    "buyer.name":      avg_field("buyer.name"),
    "buyer.tax_id":    avg_field("buyer.tax_id"),
    "amount.subtotal": avg_field("amount.subtotal"),
    "amount.vat":      avg_field("amount.vat"),
    "amount.total":    avg_field("amount.total"),
    "items ≥1":        avg_field("items ≥1"),
}
math_rate = rounds[0]["report"].get("math_consistency", {}).get("rate") or 0

# ค่าเก่า (Phase 1)
OLD = {
    "document_type":   1.000,
    "document_number": 0.399,
    "document_date":   0.379,
    "seller.name":     0.089,
    "seller.tax_id":   0.034,
    "buyer.name":      0.108,
    "buyer.tax_id":    0.296,
    "amount.subtotal": 0.995,
    "amount.vat":      0.995,
    "amount.total":    1.000,
    "items ≥1":        1.000,
}

print("\n📊 สรุปผลการทดลองที่ 3 (prompt ปรับปรุง):")
print(f"  วิเคราะห์สำเร็จ: {n_succ}/{n_total} ({succ_rate*100:.1f}%)")
print(f"\n  {'ฟิลด์':<28} {'เก่า':>7}  {'ใหม่':>7}  {'เปลี่ยน':>8}")
print(f"  {'-'*28} {'-'*7}  {'-'*7}  {'-'*8}")
for label, new_val in fields_new.items():
    old_val = OLD.get(label, 0)
    delta = new_val - old_val
    sign = "+" if delta >= 0 else ""
    flag = " ✅" if new_val >= 0.80 else (" ⚠️" if new_val >= 0.60 else " 🔴")
    print(f"  {label:<28} {old_val*100:>6.1f}%  {new_val*100:>6.1f}%  {sign}{delta*100:>5.1f}%{flag}")

print()

# ─── 2. อัปเดต docx ───────────────────────────────────────────────────────────

if not DOCX_PATH.exists():
    print(f"❌ ไม่พบ {DOCX_PATH}"); sys.exit(1)

import shutil
tmp = Path("/tmp/paper_update.docx")
shutil.copy2(DOCX_PATH, tmp)

doc = Document(tmp)
tables = doc.tables
# 0=2.1, 1=4.1, 2=4.2, 3=4.3, 4=4.4, 5=4.5, 6=4.6, 7=4.7

FONT = "TH Sarabun New"
SZ   = 32  # 16pt in half-points

def set_font(run, bold=False):
    rPr = run._r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); run._r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.insert(0, rFonts)
    for attr in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
        rFonts.set(qn(attr), FONT)
    for sn in ("w:sz","w:szCs"):
        el = rPr.find(qn(sn))
        if el is None: el = OxmlElement(sn); rPr.append(el)
        el.set(qn("w:val"), str(SZ))
    if bold:
        b = rPr.find(qn("w:b"))
        if b is None: b = OxmlElement("w:b"); rPr.append(b)

def set_cell_text(cell, text, bold=False, color=None):
    """Clear cell content and set new text with proper font."""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
    if cell.paragraphs:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()
    run = para.add_run(text)
    set_font(run, bold=bold)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def shade_cell(cell, fill_hex):
    """Add background color to cell."""
    tcPr = cell._tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr"); cell._tc.insert(0, tcPr)
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)

# ── 2.1 ตาราง 4.1: เพิ่มคอลัมน์ใหม่ ──────────────────────────────────────────

tbl_41 = tables[1]

# field_map: row index → field key
FIELD_ROW_MAP = {
    1: "document_type",
    2: "document_number",
    3: "document_date",
    4: "seller.name",
    5: "seller.tax_id",
    6: "buyer.name",
    7: "buyer.tax_id",
    8: "amount.subtotal",
    9: "amount.vat",
    10: "amount.total",
}

def add_column_to_table(tbl, header_texts, row_texts):
    """Append a column to the right of an existing table."""
    tblGrid = tbl._tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), "2000")
        tblGrid.append(gridCol)

    for i, row in enumerate(tbl.rows):
        tr = row._tr
        # Copy last cell's XML as template
        last_tc = tr.findall(qn("w:tc"))[-1]
        new_tc = copy.deepcopy(last_tc)
        # Clear text in new cell
        for p in new_tc.findall(".//" + qn("w:p")):
            for r in p.findall(qn("w:r")):
                p.remove(r)
        tr.append(new_tc)

    # Set header
    set_cell_text(tbl.rows[0].cells[-1], header_texts[0], bold=True)
    shade_cell(tbl.rows[0].cells[-1], "2E75B6")
    set_cell_text(tbl.rows[0].cells[-1], header_texts[0], bold=True)
    run = tbl.rows[0].cells[-1].paragraphs[0].runs[-1]
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Set data rows
    for i, text in enumerate(row_texts, start=1):
        if i < len(tbl.rows):
            cell = tbl.rows[i].cells[-1]
            set_cell_text(cell, text)

print("→ อัปเดตตาราง 4.1...")
new_col_data = []
for ri in range(1, 11):
    fkey = FIELD_ROW_MAP.get(ri, "")
    new_val = fields_new.get(fkey, 0)
    old_val = OLD.get(fkey, 0)
    delta = new_val - old_val
    sign = "+" if delta >= 0 else ""
    new_col_data.append(f"{new_val*100:.1f}% ({sign}{delta*100:.1f}%)")

add_column_to_table(tbl_41,
    header_texts=["การทดลองที่ 3\n(prompt ปรับปรุง)"],
    row_texts=new_col_data
)
print("  ✓ เพิ่มคอลัมน์การทดลองที่ 3 แล้ว")

# ── 2.2 ตาราง 4.6: เพิ่มแถวสรุปการทดลองที่ 3 ─────────────────────────────────

tbl_46 = tables[6]
print("→ อัปเดตตาราง 4.6...")

def add_row_to_table(tbl, cells_text):
    # Clone last row as template
    last_tr = tbl._tbl.findall(qn("w:tr"))[-1]
    new_tr = copy.deepcopy(last_tr)
    # Clear all cells
    for tc in new_tr.findall(qn("w:tc")):
        for p in tc.findall(".//" + qn("w:p")):
            for r in p.findall(qn("w:r")):
                p.remove(r)
    tbl._tbl.append(new_tr)
    row = tbl.rows[-1]
    for i, text in enumerate(cells_text):
        if i < len(row.cells):
            set_cell_text(row.cells[i], text)

doc_num_new = fields_new["document_number"]
doc_date_new = fields_new["document_date"]
best_field = f"document_type: 100%\ndocument_number: {doc_num_new*100:.1f}%\ndocument_date: {doc_date_new*100:.1f}%"

add_row_to_table(tbl_46, [
    "การทดลองที่ 3",
    "ทดสอบ prompt ที่ปรับปรุงใหม่ (pattern document_number + document_date ครบถ้วน)",
    f"204 เอกสาร (รัน {len(rounds)} รอบ)",
    "Field Accuracy (%), เปรียบเทียบ prompt เดิมและใหม่",
    best_field,
])
print("  ✓ เพิ่มแถวการทดลองที่ 3 แล้ว")

# ── 2.3 ตาราง 4.7: อัปเดตแถวปัญหา document_number / date ─────────────────────

tbl_47 = tables[7]
print("→ อัปเดตตาราง 4.7...")

if doc_num_new >= 0.80:
    # แก้แถวที่ 1 (document_number)
    row1 = tbl_47.rows[1]
    old_problem_text = row1.cells[0].text
    if "document_number" in old_problem_text or "39.9" in old_problem_text:
        set_cell_text(row1.cells[0],
            f"ความแม่นยำ document_number ปรับปรุงแล้ว ({doc_num_new*100:.1f}%)")
        set_cell_text(row1.cells[2],
            f"เพิ่มขึ้นจาก 39.9% → {doc_num_new*100:.1f}% หลังปรับ prompt")
        print(f"  ✓ อัปเดตแถว document_number: {doc_num_new*100:.1f}%")

if doc_date_new >= 0.80:
    row2 = tbl_47.rows[2]
    old_problem_text = row2.cells[0].text
    if "document_date" in old_problem_text or "37.9" in old_problem_text:
        set_cell_text(row2.cells[0],
            f"ความแม่นยำ document_date ปรับปรุงแล้ว ({doc_date_new*100:.1f}%)")
        set_cell_text(row2.cells[2],
            f"เพิ่มขึ้นจาก 37.9% → {doc_date_new*100:.1f}% หลังปรับ prompt")
        print(f"  ✓ อัปเดตแถว document_date: {doc_date_new*100:.1f}%")

# ── 2.4 บันทึกไฟล์ ────────────────────────────────────────────────────────────

doc.save(tmp)

# verify
import zipfile
with zipfile.ZipFile(tmp) as z:
    z.testzip()
print("\n✓ ตรวจสอบ zip OK")

import shutil
shutil.copy2(tmp, DOCX_PATH)
print(f"✓ บันทึก: {DOCX_PATH}")

# ── 3. สรุป ───────────────────────────────────────────────────────────────────

print("\n" + "═"*55)
print("  อัปเดต financial-ocr-paper-v2.docx เรียบร้อย")
print("  สิ่งที่เปลี่ยน:")
print("    • ตาราง 4.1 — เพิ่มคอลัมน์การทดลองที่ 3")
print("    • ตาราง 4.6 — เพิ่มแถวสรุปการทดลองที่ 3")
if doc_num_new >= 0.80 or doc_date_new >= 0.80:
    print("    • ตาราง 4.7 — อัปเดตแถวปัญหาที่แก้ไขได้แล้ว")
print("═"*55)
